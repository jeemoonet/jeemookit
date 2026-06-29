#!/usr/bin/env python3
"""Convert Markdown (with Mermaid blocks and local images) to Word (.docx)."""

from __future__ import annotations

import argparse
import hashlib
import io
import os
import re
import shutil
import subprocess
import sys
import tempfile
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable, Iterator, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from PIL import Image

SVG_RENDERER_VERSION = "puppeteer-v1"

SKILL_DIR = Path(__file__).resolve().parent.parent
PROJECT_ROOT = SKILL_DIR.parent.parent.parent

FIGURE_LINE_RE = re.compile(
    r"^(?:[-*+]\s*)?\*\*(?P<label>图[^*]+?)\*\*[（(]`(?P<path>[^`]+\.(?:svg|png|jpe?g|gif|webp))`[）)]"
    r"\s*[：:]\s*(?P<caption>.*)$",
    re.IGNORECASE,
)
# 兼容标签损坏但仍含 `assets/xxx.svg` 的行
FIGURE_PATH_LINE_RE = re.compile(
    r"^(?:[-*+]\s*)?\*\*(?P<label>[^*]+?)\*\*.*?`(?P<path>[^`]+\.(?:svg|png|jpe?g|gif|webp))`"
    r"\s*[：:]\s*(?P<caption>.*)$",
    re.IGNORECASE,
)
FIGURE_TABLE_ROW_RE = re.compile(r"^图(\d+)$")
IMAGE_MD_RE = re.compile(r"^!\[(?P<alt>[^\]]*)\]\((?P<path>[^)]+)\)\s*$")
HEADING_RE = re.compile(r"^(?P<hashes>#{1,6})\s+(?P<text>.+)$")
HR_RE = re.compile(r"^(?:-{3,}|\*{3,}|_{3,})\s*$")
LIST_UL_RE = re.compile(r"^(?P<indent>\s*)[-*+]\s+(?P<text>.+)$")
LIST_OL_RE = re.compile(r"^(?P<indent>\s*)\d+\.\s+(?P<text>.+)$")
TABLE_SEP_RE = re.compile(r"^\|?[\s\-:|]+\|?$")


@dataclass
class HeadingBlock:
    level: int
    text: str


@dataclass
class ParagraphBlock:
    text: str


@dataclass
class ListItemBlock:
    text: str
    ordered: bool
    level: int = 0


@dataclass
class TableBlock:
    rows: list[list[str]]


@dataclass
class CodeBlock:
    language: Optional[str]
    content: str


@dataclass
class MermaidBlock:
    content: str


@dataclass
class FigureBlock:
    label: str
    image_path: str
    caption: str


@dataclass
class ImageBlock:
    alt: str
    path: str


@dataclass
class BlockquoteBlock:
    text: str


@dataclass
class HRBlock:
    pass


Block = (
    HeadingBlock
    | ParagraphBlock
    | ListItemBlock
    | TableBlock
    | CodeBlock
    | MermaidBlock
    | FigureBlock
    | ImageBlock
    | BlockquoteBlock
    | HRBlock
)


def content_hash(text: str) -> str:
    return hashlib.sha256(text.encode("utf-8")).hexdigest()[:16]


def find_project_root(start: Path) -> Path:
    for candidate in [start, *start.parents]:
        if (candidate / ".cursor").is_dir():
            return candidate
    return start.parent if start.is_file() else start


def cache_dir_for(md_path: Path) -> Path:
    root = find_project_root(md_path.resolve())
    cache = root / ".cache" / "md2docx"
    cache.mkdir(parents=True, exist_ok=True)
    return cache


def resolve_image_path(md_dir: Path, rel_path: str) -> Path:
    return resolve_figure_path(md_dir, rel_path)


def normalize_figure_filename(name: str) -> str:
    name = name.strip().replace("\\", "/")
    if "/" in name:
        name = name.rsplit("/", 1)[-1]
    match = re.match(r"^([^\s（(]+)", name)
    return match.group(1) if match else name


def figure_sort_key(name: str) -> tuple[int, str]:
    match = re.search(r"(?:图|fig)(\d+)", name, re.IGNORECASE)
    return (int(match.group(1)), name.lower()) if match else (999, name.lower())


def resolve_figure_path(md_dir: Path, rel_path: str) -> Path:
    rel_path = rel_path.strip().replace("\\", "/")
    if rel_path.startswith("<") and rel_path.endswith(">"):
        rel_path = rel_path[1:-1]
    filename = normalize_figure_filename(rel_path)

    candidates = [
        (md_dir / rel_path).resolve(),
        (md_dir / "assets" / filename).resolve(),
        (md_dir / "图例" / filename).resolve(),  # 旧路径兼容
        (md_dir / filename).resolve(),
    ]
    seen: set[Path] = set()
    for candidate in candidates:
        if candidate in seen:
            continue
        seen.add(candidate)
        if candidate.is_file():
            return candidate

    for subdir in ("assets", "图例"):
        fig_dir = md_dir / subdir
        if not fig_dir.is_dir():
            continue
        number_match = re.match(r"(?:fig|(图))(\d+)", filename, re.IGNORECASE)
        if number_match:
            number = number_match.group(2)
            for pattern in (f"图{number}*.svg", f"fig{number}.svg", f"Fig{number}.svg"):
                matches = sorted(fig_dir.glob(pattern), key=lambda p: figure_sort_key(p.name))
                if matches:
                    return matches[0].resolve()

    raise FileNotFoundError(f"图片不存在: {rel_path}（相对 {md_dir}）")


def is_figure_table(table: TableBlock) -> bool:
    if not table.rows:
        return False
    header = [cell.strip() for cell in table.rows[0]]
    if len(header) < 3:
        return False
    joined = "".join(header)
    return "图号" in joined and "文件名" in joined


def figure_blocks_from_table(table: TableBlock) -> list[FigureBlock]:
    figures: list[FigureBlock] = []
    for row in table.rows[1:]:
        if len(row) < 3:
            continue
        label, filename, caption = row[0].strip(), row[1].strip(), row[2].strip()
        if not FIGURE_TABLE_ROW_RE.match(label):
            continue
        rel_path = f"assets/{normalize_figure_filename(filename)}"
        figures.append(FigureBlock(label=label, image_path=rel_path, caption=caption))
    return figures


def parse_legend_description(md_dir: Path) -> list[tuple[str, str, str]]:
    legend = md_dir / "图例说明.md"
    if not legend.is_file():
        return []

    rows: list[tuple[str, str, str]] = []
    for line in legend.read_text(encoding="utf-8").splitlines():
        stripped = line.strip()
        if not stripped.startswith("|") or TABLE_SEP_RE.match(stripped.replace("|", "").strip()):
            continue
        cells = [cell.strip() for cell in stripped.strip("|").split("|")]
        if len(cells) < 3 or cells[0] == "图号":
            continue
        if FIGURE_TABLE_ROW_RE.match(cells[0]):
            rows.append((cells[0], cells[1], cells[2]))
    return rows


def scan_figure_directory(md_dir: Path) -> list[tuple[str, str, str]]:
    rows: list[tuple[str, str, str]] = []
    for subdir in ("assets", "图例"):
        fig_dir = md_dir / subdir
        if not fig_dir.is_dir():
            continue
        for svg_path in sorted(fig_dir.glob("*.svg"), key=lambda p: figure_sort_key(p.name)):
            label_match = re.match(r"^(图\d+|fig\d+)", svg_path.name, re.IGNORECASE)
            if not label_match:
                continue
            raw_label = label_match.group(1)
            label = f"图{raw_label[3:]}" if raw_label.lower().startswith("fig") else raw_label
            rows.append((label, svg_path.name, ""))
    return rows


def supplement_figure_blocks(blocks: list[Block], md_dir: Path) -> list[Block]:
    result: list[Block] = []
    embedded_labels: set[str] = set()

    for block in blocks:
        result.append(block)
        if isinstance(block, FigureBlock):
            embedded_labels.add(block.label)
        elif isinstance(block, TableBlock) and is_figure_table(block):
            for figure in figure_blocks_from_table(block):
                if figure.label in embedded_labels:
                    continue
                result.append(figure)
                embedded_labels.add(figure.label)

    sources = parse_legend_description(md_dir) or scan_figure_directory(md_dir)
    to_insert: list[FigureBlock] = []
    for label, filename, caption in sources:
        if label in embedded_labels:
            continue
        rel_path = f"assets/{normalize_figure_filename(filename)}"
        to_insert.append(FigureBlock(label=label, image_path=rel_path, caption=caption))
        embedded_labels.add(label)

    if not to_insert:
        return result

    insert_at = len(result)
    for index, block in enumerate(result):
        if isinstance(block, HeadingBlock) and block.level == 2:
            text = block.text.strip()
            if text.startswith("6") or text.startswith("6."):
                insert_at = index
                break

    for offset, figure in enumerate(to_insert):
        result.insert(insert_at + offset, figure)
    return result


def find_node() -> str:
    node = shutil.which("node") or shutil.which("node.exe")
    if not node:
        raise RuntimeError("未找到 Node.js，无法渲染 SVG 中文文本")
    return node


def svg_to_png(svg_path: Path, png_path: Path) -> None:
    """Render SVG via headless Chromium (svglib/reportlab cannot render CJK text)."""
    node = find_node()
    script = SKILL_DIR / "scripts" / "svg2png.mjs"
    if not script.is_file():
        raise RuntimeError(f"缺少 SVG 渲染脚本: {script}")

    png_path.parent.mkdir(parents=True, exist_ok=True)
    cmd = [node, str(script), str(svg_path.resolve()), str(png_path.resolve())]
    result = subprocess.run(cmd, capture_output=True, text=True, cwd=str(SKILL_DIR))
    if result.returncode != 0 or not png_path.is_file():
        detail = (result.stderr or result.stdout or "未知错误").strip()
        raise RuntimeError(f"SVG 渲染失败 ({svg_path.name}): {detail}")


def raster_to_png(image_path: Path, png_path: Path) -> None:
    png_path.parent.mkdir(parents=True, exist_ok=True)
    if image_path.suffix.lower() == ".png":
        shutil.copy2(image_path, png_path)
        return
    with Image.open(image_path) as img:
        if img.mode in ("RGBA", "LA", "P"):
            img = img.convert("RGBA")
        else:
            img = img.convert("RGB")
        img.save(png_path, format="PNG")


def ensure_png(image_path: Path, cache: Path) -> Path:
    key = content_hash(
        f"{SVG_RENDERER_VERSION}:{image_path.resolve()}:{image_path.stat().st_mtime_ns}"
    )
    png_path = cache / f"img-{key}.png"
    if png_path.is_file():
        return png_path
    if image_path.suffix.lower() == ".svg":
        svg_to_png(image_path, png_path)
    else:
        raster_to_png(image_path, png_path)
    return png_path


def find_mmdc() -> list[str]:
    local = SKILL_DIR / "node_modules" / ".bin"
    if os.name == "nt":
        candidates = [local / "mmdc.cmd", local / "mmdc"]
    else:
        candidates = [local / "mmdc"]
    for exe in candidates:
        if exe.is_file():
            return [str(exe)]
    npx = shutil.which("npx") or shutil.which("npx.cmd")
    if not npx:
        raise RuntimeError("未找到 npx，请安装 Node.js 并在 skill 目录运行 npm install")
    return [npx, "--yes", "@mermaid-js/mermaid-cli"]


def render_mermaid(content: str, cache: Path) -> Path:
    key = content_hash(content)
    png_path = cache / f"mermaid-{key}.png"
    if png_path.is_file():
        return png_path

    mmdc_cmd = find_mmdc()
    with tempfile.TemporaryDirectory(prefix="md2docx-mermaid-") as tmp:
        tmp_dir = Path(tmp)
        mmd_file = tmp_dir / f"{key}.mmd"
        out_file = tmp_dir / f"{key}.png"
        mmd_file.write_text(content.strip() + "\n", encoding="utf-8")

        cmd = [
            *mmdc_cmd,
            "-i",
            str(mmd_file),
            "-o",
            str(out_file),
            "-b",
            "white",
            "-w",
            "1200",
            "-s",
            "2",
        ]
        result = subprocess.run(cmd, capture_output=True, text=True, cwd=str(SKILL_DIR))
        if result.returncode != 0 or not out_file.is_file():
            detail = (result.stderr or result.stdout or "未知错误").strip()
            raise RuntimeError(f"Mermaid 渲染失败: {detail}")

        png_path.parent.mkdir(parents=True, exist_ok=True)
        shutil.copy2(out_file, png_path)
    return png_path


def parse_table_lines(lines: list[str]) -> TableBlock:
    rows: list[list[str]] = []
    for line in lines:
        stripped = line.strip()
        if not stripped.startswith("|"):
            continue
        if TABLE_SEP_RE.match(stripped.replace("|", "").strip()):
            continue
        cells = [cell.strip() for cell in stripped.strip("|").split("|")]
        rows.append(cells)
    return TableBlock(rows=rows)


def indent_to_level(indent: str) -> int:
    """Markdown list indent: 2 spaces (or 1 tab) per nesting level."""
    expanded = indent.replace("\t", "  ")
    return min(len(expanded) // 2, 8)


def list_style_name(ordered: bool, level: int) -> str:
    base = "List Number" if ordered else "List Bullet"
    if level <= 0:
        return base
    return f"{base} {level + 1}"


def parse_blocks(text: str) -> list[Block]:
    lines = text.splitlines()
    blocks: list[Block] = []
    i = 0
    while i < len(lines):
        line = lines[i]

        if line.strip().startswith("```"):
            fence = line.strip()
            lang = fence[3:].strip() or None
            i += 1
            code_lines: list[str] = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            content = "\n".join(code_lines)
            if lang and lang.lower() == "mermaid":
                blocks.append(MermaidBlock(content=content))
            else:
                blocks.append(CodeBlock(language=lang, content=content))
            if i < len(lines):
                i += 1
            continue

        heading = HEADING_RE.match(line)
        if heading:
            blocks.append(
                HeadingBlock(level=len(heading.group("hashes")), text=heading.group("text"))
            )
            i += 1
            continue

        if HR_RE.match(line.strip()):
            blocks.append(HRBlock())
            i += 1
            continue

        if line.strip().startswith("|"):
            table_lines: list[str] = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            blocks.append(parse_table_lines(table_lines))
            continue

        figure = FIGURE_LINE_RE.match(line.strip())
        if not figure:
            figure = FIGURE_PATH_LINE_RE.match(line.strip())
        if figure:
            blocks.append(
                FigureBlock(
                    label=figure.group("label"),
                    image_path=figure.group("path"),
                    caption=figure.group("caption"),
                )
            )
            i += 1
            continue

        image = IMAGE_MD_RE.match(line.strip())
        if image:
            blocks.append(ImageBlock(alt=image.group("alt"), path=image.group("path")))
            i += 1
            continue

        if line.startswith(">"):
            quote_lines: list[str] = []
            while i < len(lines) and lines[i].startswith(">"):
                quote_lines.append(lines[i].lstrip(">").lstrip())
                i += 1
            blocks.append(BlockquoteBlock(text="\n".join(quote_lines)))
            continue

        ul = LIST_UL_RE.match(line)
        if ul:
            blocks.append(
                ListItemBlock(
                    text=ul.group("text"),
                    ordered=False,
                    level=indent_to_level(ul.group("indent")),
                )
            )
            i += 1
            continue

        ol = LIST_OL_RE.match(line)
        if ol:
            blocks.append(
                ListItemBlock(
                    text=ol.group("text"),
                    ordered=True,
                    level=indent_to_level(ol.group("indent")),
                )
            )
            i += 1
            continue

        if not line.strip():
            i += 1
            continue

        para_lines = [line]
        i += 1
        while i < len(lines):
            nxt = lines[i]
            if not nxt.strip():
                break
            if nxt.strip().startswith("```"):
                break
            if HEADING_RE.match(nxt):
                break
            if HR_RE.match(nxt.strip()):
                break
            if nxt.strip().startswith("|"):
                break
            if FIGURE_LINE_RE.match(nxt.strip()):
                break
            if IMAGE_MD_RE.match(nxt.strip()):
                break
            if nxt.startswith(">"):
                break
            if LIST_UL_RE.match(nxt) or LIST_OL_RE.match(nxt):
                break
            para_lines.append(nxt)
            i += 1
        blocks.append(ParagraphBlock(text="\n".join(para_lines)))

    return blocks


HEADING_FONT_LATIN = "Microsoft YaHei"
HEADING_FONT_CJK = "微软雅黑"

HEADING_RUN_SPECS = {
    1: {"size": 22, "color": RGBColor(0x1A, 0x1A, 0x2E)},
    2: {"size": 16, "color": RGBColor(0x2C, 0x3E, 0x50)},
    3: {"size": 14, "color": RGBColor(0x34, 0x49, 0x5E)},
    4: {"size": 12, "color": RGBColor(0x34, 0x49, 0x5E)},
    5: {"size": 11, "color": RGBColor(0x44, 0x55, 0x66)},
    6: {"size": 11, "color": RGBColor(0x55, 0x66, 0x77), "italic": True},
}


def set_run_font(
    run,
    size_pt: Optional[float] = None,
    bold: Optional[bool] = None,
    mono: bool = False,
    font_name: Optional[str] = None,
    east_asia: Optional[str] = None,
) -> None:
    if mono:
        font_name = font_name or "Consolas"
        east_asia = east_asia or "Microsoft YaHei"
    else:
        font_name = font_name or "宋体"
        east_asia = east_asia or "宋体"

    run.font.name = font_name
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:ascii"), font_name)
    r_fonts.set(qn("w:hAnsi"), font_name)
    r_fonts.set(qn("w:eastAsia"), east_asia)
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bold


def set_run_shading(run, fill_color: str) -> None:
    r_pr = run._element.get_or_add_rPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_color)
    r_pr.append(shd)


def set_paragraph_shading(paragraph, fill_color: str) -> None:
    p_pr = paragraph._element.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_color)
    p_pr.append(shd)


def set_paragraph_left_border(paragraph, color: str = "AEB5BC", width: str = "24") -> None:
    """Markdown-style blockquote: thick left bar only."""
    p_pr = paragraph._element.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), width)
    left.set(qn("w:space"), "10")
    left.set(qn("w:color"), color)
    p_bdr.append(left)
    p_pr.append(p_bdr)


def add_blockquote_runs(paragraph, text: str) -> None:
    """Inline formatting inside blockquote with muted body text."""
    quote_color = RGBColor(0x57, 0x60, 0x6A)
    pattern = re.compile(r"(\*\*.+?\*\*|`[^`]+`)")
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos : match.start()])
            set_run_font(run, size_pt=10.5)
            run.font.color.rgb = quote_color
        token = match.group(0)
        if token.startswith("**") and token.endswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size_pt=10.5, bold=True)
            run.font.color.rgb = quote_color
        elif token.startswith("`") and token.endswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, size_pt=10, mono=True)
            run.font.color.rgb = RGBColor(0x09, 0x69, 0xDA)
            set_run_shading(run, "EBEEF2")
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_run_font(run, size_pt=10.5)
        run.font.color.rgb = quote_color


def add_blockquote(document: Document, text: str) -> None:
    """Render Markdown blockquote: left gray bar + light background + muted text."""
    lines = [line.strip() for line in text.split("\n") if line.strip()]
    if not lines:
        return

    for idx, line in enumerate(lines):
        paragraph = document.add_paragraph()
        is_first = idx == 0
        is_last = idx == len(lines) - 1

        set_paragraph_left_border(paragraph)
        set_paragraph_shading(paragraph, "F6F8FA")
        paragraph.paragraph_format.left_indent = Cm(0.55)
        paragraph.paragraph_format.space_before = Pt(4 if is_first else 0)
        paragraph.paragraph_format.space_after = Pt(4 if is_last else 0)
        add_blockquote_runs(paragraph, line)


def add_horizontal_rule(document: Document) -> None:
    """Render Markdown horizontal rule as a thin separator line."""
    paragraph = document.add_paragraph()
    p_pr = paragraph._element.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "D0D7DE")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)
    paragraph.paragraph_format.space_before = Pt(10)
    paragraph.paragraph_format.space_after = Pt(10)


def set_cell_shading(cell, fill_color: str) -> None:
    tc_pr = cell._element.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_color)
    tc_pr.append(shd)


def set_cell_border(cell, color: str = "D0D7DE") -> None:
    tc_pr = cell._element.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), color)
        tc_borders.append(border)
    tc_pr.append(tc_borders)


def clear_table_borders(table) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl.insert(0, tbl_pr)
    tbl_borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "none")
        border.set(qn("w:sz"), "0")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "auto")
        tbl_borders.append(border)
    tbl_pr.append(tbl_borders)


def add_inline_runs(paragraph, text: str, size_pt: float = 11, mono: bool = False) -> None:
    pattern = re.compile(r"(\*\*.+?\*\*|`[^`]+`)")
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos : match.start()])
            set_run_font(run, size_pt=size_pt, mono=mono)
        token = match.group(0)
        if token.startswith("**") and token.endswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size_pt=size_pt, bold=True, mono=mono)
        elif token.startswith("`") and token.endswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, size_pt=max(size_pt - 0.5, 9), mono=True)
            run.font.color.rgb = RGBColor(0x09, 0x69, 0xDA)
            set_run_shading(run, "F6F8FA")
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_run_font(run, size_pt=size_pt, mono=mono)


def add_paragraph(document: Document, text: str, style: Optional[str] = None) -> None:
    paragraph = document.add_paragraph(style=style)
    add_inline_runs(paragraph, text)


def add_list_item(document: Document, text: str, ordered: bool, level: int = 0) -> None:
    """Render Markdown list item with Word native bullet/numbering styles."""
    style_name = list_style_name(ordered, level)
    try:
        paragraph = document.add_paragraph(style=style_name)
    except KeyError:
        paragraph = document.add_paragraph()
        prefix = "  " * level + ("1. " if ordered else "• ")
        add_inline_runs(paragraph, prefix + text)
        paragraph.paragraph_format.left_indent = Cm(0.74 * (level + 1))
        paragraph.paragraph_format.first_line_indent = Cm(-0.37)
        return

    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.line_spacing = 1.15
    add_inline_runs(paragraph, text)


def set_style_font(style, font_name: str, east_asia: Optional[str] = None) -> None:
    east_asia = east_asia or font_name
    style.font.name = font_name
    r_pr = style.element.rPr
    if r_pr is None:
        r_pr = OxmlElement("w:rPr")
        style.element.append(r_pr)
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:ascii"), font_name)
    r_fonts.set(qn("w:hAnsi"), font_name)
    r_fonts.set(qn("w:eastAsia"), east_asia)


def configure_heading_styles(document: Document) -> None:
    """Apply heading styles (Word Heading 1–6), Feishu-friendly 微软雅黑 + outline level."""
    specs = {
        1: {
            "size": 22,
            "font": "微软雅黑",
            "color": RGBColor(0x1A, 0x1A, 0x2E),
            "space_before": 24,
            "space_after": 12,
            "center": True,
        },
        2: {
            "size": 16,
            "font": "微软雅黑",
            "color": RGBColor(0x2C, 0x3E, 0x50),
            "space_before": 18,
            "space_after": 8,
        },
        3: {
            "size": 14,
            "font": "微软雅黑",
            "color": RGBColor(0x34, 0x49, 0x5E),
            "space_before": 14,
            "space_after": 6,
        },
        4: {
            "size": 12,
            "font": "微软雅黑",
            "color": RGBColor(0x34, 0x49, 0x5E),
            "space_before": 10,
            "space_after": 4,
        },
        5: {
            "size": 11,
            "font": "微软雅黑",
            "color": RGBColor(0x44, 0x55, 0x66),
            "space_before": 8,
            "space_after": 4,
        },
        6: {
            "size": 11,
            "font": "微软雅黑",
            "color": RGBColor(0x55, 0x66, 0x77),
            "space_before": 6,
            "space_after": 2,
            "italic": True,
        },
    }
    for level, spec in specs.items():
        style = document.styles[f"Heading {level}"]
        set_style_font(style, HEADING_FONT_LATIN, spec["font"])
        style.font.size = Pt(spec["size"])
        style.font.bold = True
        style.font.color.rgb = spec["color"]
        if spec.get("italic"):
            style.font.italic = True
        pf = style.paragraph_format
        pf.space_before = Pt(spec["space_before"])
        pf.space_after = Pt(spec["space_after"])
        pf.keep_with_next = True
        pf.outline_level = level - 1
        pf.line_spacing = 1.2
        if spec.get("center"):
            pf.alignment = WD_ALIGN_PARAGRAPH.CENTER


def configure_list_styles(document: Document) -> None:
    """Configure built-in list styles for nested bullets/numbers (Feishu import friendly)."""
    for name in (
        "List Bullet",
        "List Bullet 2",
        "List Bullet 3",
        "List Bullet 4",
        "List Bullet 5",
        "List Number",
        "List Number 2",
        "List Number 3",
        "List Number 4",
        "List Number 5",
    ):
        try:
            style = document.styles[name]
        except KeyError:
            continue
        suffix = name.rsplit(" ", 1)[-1]
        level = 0 if suffix in ("Bullet", "Number") else max(int(suffix) - 1, 0)
        pf = style.paragraph_format
        pf.left_indent = Cm(0.74 * (level + 1))
        pf.first_line_indent = Cm(-0.37)
        pf.space_before = Pt(0)
        pf.space_after = Pt(3)
        pf.line_spacing = 1.15
        set_style_font(style, "宋体", "宋体")
        style.font.size = Pt(11)


def set_paragraph_outline_level(paragraph, level: int) -> None:
    p_pr = paragraph._element.get_or_add_pPr()
    outline = p_pr.find(qn("w:outlineLvl"))
    if outline is None:
        outline = OxmlElement("w:outlineLvl")
        p_pr.append(outline)
    outline.set(qn("w:val"), str(max(0, min(level - 1, 8))))


def add_heading_runs(paragraph, text: str, level: int) -> None:
    """Add inline formatting inside a heading with explicit 黑体 run fonts."""
    spec = HEADING_RUN_SPECS.get(level, HEADING_RUN_SPECS[2])
    size_pt = spec["size"]
    color = spec["color"]
    italic = spec.get("italic", False)

    def apply_heading_font(run, *, bold: bool = True, mono: bool = False) -> None:
        if mono:
            set_run_font(run, size_pt=max(size_pt - 1, 9), mono=True)
        else:
            set_run_font(
                run,
                size_pt=size_pt,
                bold=bold,
                font_name=HEADING_FONT_LATIN,
                east_asia=HEADING_FONT_CJK,
            )
            run.font.color.rgb = color
            run.italic = italic

    pattern = re.compile(r"(\*\*.+?\*\*|`[^`]+`)")
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos : match.start()])
            apply_heading_font(run)
        token = match.group(0)
        if token.startswith("**") and token.endswith("**"):
            run = paragraph.add_run(token[2:-2])
            apply_heading_font(run, bold=True)
        elif token.startswith("`") and token.endswith("`"):
            run = paragraph.add_run(token[1:-1])
            apply_heading_font(run, bold=False, mono=True)
            run.font.color.rgb = RGBColor(0x09, 0x69, 0xDA)
            set_run_shading(run, "F6F8FA")
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        apply_heading_font(run)


def add_heading(document: Document, level: int, text: str) -> None:
    level = max(1, min(level, 6))
    paragraph = document.add_paragraph(style=f"Heading {level}")
    set_paragraph_outline_level(paragraph, level)
    add_heading_runs(paragraph, text, level)


def plain_cell_text(text: str) -> str:
    """Strip Markdown markers for width estimation."""
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"`([^`]+)`", r"\1", text)
    return text.strip()


def text_display_width(text: str) -> float:
    """Estimate visual width: CJK/full-width chars count as 2, others as 1."""
    width = 0.0
    for ch in text:
        if ord(ch) > 127:
            width += 2.0
        else:
            width += 1.0
    return width


def set_cell_no_wrap(cell, nowrap: bool) -> None:
    tc_pr = cell._element.get_or_add_tcPr()
    no_wrap_el = tc_pr.find(qn("w:noWrap"))
    if nowrap:
        if no_wrap_el is None:
            tc_pr.append(OxmlElement("w:noWrap"))
    elif no_wrap_el is not None:
        tc_pr.remove(no_wrap_el)


def configure_table_autofit(table) -> None:
    """Auto-fit table/column width to content (maps better to Feishu docx import)."""
    table.autofit = True
    table.allow_autofit = True

    tbl = table._tbl
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl.insert(0, tbl_pr)

    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.insert(0, tbl_layout)
    tbl_layout.set(qn("w:type"), "autofit")

    for tag in ("w:tblW",):
        existing = tbl_pr.find(qn(tag))
        if existing is not None:
            tbl_pr.remove(existing)
    tbl_w = OxmlElement("w:tblW")
    tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "pct")
    tbl_w.set(qn("w:w"), "5000")

    tbl_grid = tbl.tblGrid
    if tbl_grid is not None:
        tbl.remove(tbl_grid)

    for row in table.rows:
        for cell in row.cells:
            tc_pr = cell._element.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.insert(0, tc_w)
            tc_w.set(qn("w:type"), "auto")
            tc_w.set(qn("w:w"), "0")
            set_cell_no_wrap(cell, False)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(2)
                paragraph.paragraph_format.space_after = Pt(2)


def compute_column_widths(rows: list[list[str]], col_count: int, max_width_cm: float = 16.5) -> list[Cm]:
    """Distribute table width proportionally by column content."""
    max_units = [1.0] * col_count
    for row in rows:
        for c_idx in range(col_count):
            cell_text = row[c_idx] if c_idx < len(row) else ""
            units = text_display_width(plain_cell_text(cell_text))
            max_units[c_idx] = max(max_units[c_idx], units)

    total_units = sum(max_units) or 1.0
    min_cm = 1.6
    raw_cm = [max(min_cm, max_width_cm * units / total_units) for units in max_units]
    scale = max_width_cm / sum(raw_cm)
    return [Cm(w * scale) for w in raw_cm]


def set_table_column_widths(table, widths: list[Cm]) -> None:
    tbl = table._tbl
    tbl_grid = tbl.tblGrid
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        tbl.insert(0, tbl_grid)
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width.twips))
        tbl_grid.append(grid_col)
    for row in table.rows:
        for c_idx, cell in enumerate(row.cells):
            if c_idx < len(widths):
                cell.width = widths[c_idx]


def add_table_cell_runs(paragraph, text: str, *, header: bool = False) -> None:
    add_inline_runs(paragraph, text, size_pt=10)
    if header:
        for run in paragraph.runs:
            run.bold = True


def add_table(document: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    col_count = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=col_count)
    table.style = "Table Grid"
    configure_table_autofit(table)

    for r_idx, row in enumerate(rows):
        is_header = r_idx == 0
        for c_idx in range(col_count):
            cell_text = row[c_idx] if c_idx < len(row) else ""
            cell = table.rows[r_idx].cells[c_idx]
            cell.text = ""
            paragraph = cell.paragraphs[0]
            add_table_cell_runs(paragraph, cell_text, header=is_header)
            if is_header:
                set_cell_shading(cell, "F0F3F6")
                for run in paragraph.runs:
                    set_run_font(run, size_pt=10, bold=True, font_name="微软雅黑", east_asia="微软雅黑")


def image_display_width(png_path: Path, max_cm: float = 15.0) -> Cm:
    with Image.open(png_path) as img:
        width_px, _ = img.size
    if width_px <= 0:
        return Cm(max_cm)
    # Rough conversion assuming ~96 DPI rendered assets
    width_cm = min(max_cm, max(8.0, width_px / 96.0 * 2.54))
    return Cm(width_cm)


def add_centered_image(document: Document, png_path: Path) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(png_path), width=image_display_width(png_path))


def add_code_block(document: Document, language: Optional[str], content: str) -> None:
    """Render fenced code block as a shaded, bordered box (GitHub-like light theme)."""
    table = document.add_table(rows=1, cols=1)
    table.autofit = True
    clear_table_borders(table)
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, "F6F8FA")
    set_cell_border(cell, "D0D7DE")

    cell.text = ""
    lines = content.split("\n")
    if language:
        lang_p = cell.paragraphs[0]
        lang_run = lang_p.add_run(language)
        set_run_font(lang_run, size_pt=8, mono=True)
        lang_run.font.color.rgb = RGBColor(0x65, 0x6D, 0x76)
        lang_run.italic = True

    for idx, line in enumerate(lines):
        paragraph = cell.add_paragraph() if (language or idx > 0) else cell.paragraphs[0]
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.15
        run = paragraph.add_run(line if line else " ")
        set_run_font(run, size_pt=9, mono=True)
        run.font.color.rgb = RGBColor(0x24, 0x29, 0x2E)

    spacer = document.add_paragraph()
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after = Pt(6)


def embed_image(document: Document, md_dir: Path, rel_path: str, cache: Path, errors: list[str]) -> None:
    try:
        image_path = resolve_figure_path(md_dir, rel_path)
        png_path = ensure_png(image_path, cache)
        add_centered_image(document, png_path)
    except Exception as exc:  # noqa: BLE001 - collect and continue
        errors.append(f"{rel_path}: {exc}")


def blocks_to_docx(blocks: list[Block], md_path: Path, output_path: Path) -> list[str]:
    md_dir = md_path.parent
    cache = cache_dir_for(md_path)
    document = Document()
    errors: list[str] = []

    normal = document.styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(11)
    configure_heading_styles(document)
    configure_list_styles(document)

    for block in blocks:
        if isinstance(block, HeadingBlock):
            add_heading(document, block.level, block.text)
        elif isinstance(block, ParagraphBlock):
            add_paragraph(document, block.text)
        elif isinstance(block, ListItemBlock):
            add_list_item(document, block.text, block.ordered, block.level)
        elif isinstance(block, TableBlock):
            add_table(document, block.rows)
        elif isinstance(block, CodeBlock):
            add_code_block(document, block.language, block.content)
        elif isinstance(block, MermaidBlock):
            try:
                png_path = render_mermaid(block.content, cache)
                add_centered_image(document, png_path)
            except Exception as exc:  # noqa: BLE001
                errors.append(f"Mermaid: {exc}")
                add_code_block(document, "mermaid", block.content)
        elif isinstance(block, FigureBlock):
            caption = f"**{block.label}**：{block.caption}"
            add_paragraph(document, caption)
            embed_image(document, md_dir, block.image_path, cache, errors)
        elif isinstance(block, ImageBlock):
            if block.alt:
                add_paragraph(document, block.alt)
            embed_image(document, md_dir, block.path, cache, errors)
        elif isinstance(block, BlockquoteBlock):
            add_blockquote(document, block.text)
        elif isinstance(block, HRBlock):
            add_horizontal_rule(document)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_path))
    return errors


def convert(md_path: Path, output_path: Optional[Path] = None) -> Path:
    md_path = md_path.resolve()
    if not md_path.is_file():
        raise FileNotFoundError(f"Markdown 文件不存在: {md_path}")
    if output_path is None:
        output_path = md_path.with_suffix(".docx")
    else:
        output_path = output_path.resolve()

    text = md_path.read_text(encoding="utf-8")
    blocks = parse_blocks(text)
    blocks = supplement_figure_blocks(blocks, md_path.parent)
    errors = blocks_to_docx(blocks, md_path, output_path)

    if errors:
        print("导出完成，但以下资源处理失败：", file=sys.stderr)
        for err in errors:
            print(f"  - {err}", file=sys.stderr)

    print(f"已生成: {output_path}")
    return output_path


def main(argv: Optional[list[str]] = None) -> int:
    parser = argparse.ArgumentParser(description="Markdown（含 Mermaid/图片）转 Word")
    parser.add_argument("markdown", type=Path, help="输入 .md 文件路径")
    parser.add_argument("-o", "--output", type=Path, help="输出 .docx 路径（默认同名）")
    args = parser.parse_args(argv)

    try:
        convert(args.markdown, args.output)
    except Exception as exc:  # noqa: BLE001
        print(f"错误: {exc}", file=sys.stderr)
        return 1
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
